# utilities.py

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
import os
from data_handling import load_and_extract_text, extract_applicant_name
from docx import Document
from datetime import datetime
import os

from reportlab.lib.pagesizes import letter
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, PageBreak
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
import os
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor, Inches
from datetime import datetime
import os

def remove_control_characters(s):
    return re.sub(r'[\x00-\x1F\x7F-\x9F]', '', s)


def create_pdf(output_pdf, sender_name, cover_letter_text):
    """
    Create a PDF cover letter from the provided text.

    Args:
        output_pdf (str): The output PDF file path.
        sender_name (str): The name of the sender.
        cover_letter_text (str): The full text of the cover letter.
    """
    # Ensure the output directory exists
    os.makedirs(os.path.dirname(output_pdf), exist_ok=True)

    # Create the PDF document
    pdf = SimpleDocTemplate(
        output_pdf, pagesize=letter,
        rightMargin=72, leftMargin=72,
        topMargin=72, bottomMargin=18
    )
    styles = getSampleStyleSheet()
    story = []

    # Custom styles
    body_style = ParagraphStyle(
        name='BodyStyle', fontSize=12,
        leading=14, spaceAfter=12, alignment=0
    )
    closing_style = ParagraphStyle(
        name='ClosingStyle', fontSize=12,
        leading=14, spaceAfter=6, alignment=0
    )
    name_style = ParagraphStyle(
        name='NameStyle', fontSize=12,
        leading=14, spaceAfter=6, alignment=0,
        fontName='Helvetica-Bold'
    )

    # Split the cover letter into paragraphs
    paragraphs = cover_letter_text.strip().split('\n\n')
    for para in paragraphs:
        story.append(Paragraph(para.strip(), body_style))
        story.append(Spacer(1, 12))

    # Add closing and sender's name
    story.append(Paragraph('Best regards,', closing_style))
    story.append(Spacer(1, 24))
    story.append(Paragraph(sender_name, name_style))

    # Build the PDF
    pdf.build(story)


def create_cover_letter(text, file_name='Cover_Letter.docx'):
    """
    Generates a .docx file from the input text, formatted as a professional cover letter.

    Parameters:
    text (str): The body of the cover letter.
    file_name (str): The desired file name for the saved document (default: 'Cover_Letter.docx').

    Returns:
    str: Path to the generated .docx file.
    """

    # Split the text into lines at double newlines for paragraphs
    paragraphs = text.split("\n\n")

    # Create a new Document
    doc = Document()

    # Add each paragraph to the document
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)

    # Save the document with the given file name
    doc.save(file_name)

    return file_name


def create_docx(output_docx_path, sender_name, cover_letter):
    """
    Creates a Word document for a cover letter with enhanced formatting.

    Args:
        output_docx_path (str): The file path where the Word document will be saved.
        sender_name (str): The name of the sender.
        cover_letter (str): The body of the cover letter.
    """
    # Create a new Document
    doc = Document()

    # Set up the document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Define custom styles
    styles = doc.styles

    # Style for sender's name
    name_style = styles.add_style('NameStyle', WD_STYLE_TYPE.PARAGRAPH)
    name_style.font.name = 'Calibri'
    name_style.font.size = Pt(14)
    name_style.font.bold = True

    # Style for contact information
    contact_style = styles.add_style('ContactStyle', WD_STYLE_TYPE.PARAGRAPH)
    contact_style.font.name = 'Calibri'
    contact_style.font.size = Pt(11)

    # Style for date
    date_style = styles.add_style('DateStyle', WD_STYLE_TYPE.PARAGRAPH)
    date_style.font.name = 'Calibri'
    date_style.font.size = Pt(11)
    date_style.font.italic = True

    # Style for recipient info and body text
    normal_style = styles.add_style('NormalStyle', WD_STYLE_TYPE.PARAGRAPH)
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)

    # Style for closing
    closing_style = styles.add_style('ClosingStyle', WD_STYLE_TYPE.PARAGRAPH)
    closing_style.font.name = 'Calibri'
    closing_style.font.size = Pt(11)

    # Add Sender's Name and Contact Information
    p = doc.add_paragraph(sender_name, style='NameStyle')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph('Your Address Line 1', style='ContactStyle')
    doc.add_paragraph('Your Address Line 2', style='ContactStyle')
    doc.add_paragraph('Phone: Your Phone Number', style='ContactStyle')
    doc.add_paragraph('Email: your.email@example.com', style='ContactStyle')
    doc.add_paragraph('')  # Empty paragraph for spacing

    # Add Date
    date_paragraph = doc.add_paragraph(datetime.now().strftime('%B %d, %Y'), style='DateStyle')
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph('')  # Empty paragraph for spacing

    # Add Recipient's Name and Address
    doc.add_paragraph('Hiring Manager', style='NormalStyle')
    doc.add_paragraph('Company Name', style='NormalStyle')
    doc.add_paragraph('Company Address Line 1', style='NormalStyle')
    doc.add_paragraph('Company Address Line 2', style='NormalStyle')
    doc.add_paragraph('')  # Empty paragraph for spacing

    # Add Salutation
    doc.add_paragraph('Dear Hiring Manager,', style='NormalStyle')
    doc.add_paragraph('')  # Empty paragraph for spacing

    # Add Body Paragraphs
    # Split the cover letter text into paragraphs based on double newlines
    paragraphs = cover_letter.strip().split('\n\n')
    for paragraph in paragraphs:
        p = doc.add_paragraph(paragraph.strip(), style='NormalStyle')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(12)

    # Add Closing
    doc.add_paragraph('Best regards,', style='ClosingStyle')
    doc.add_paragraph('')  # Empty paragraph for spacing
    doc.add_paragraph(sender_name, style='NameStyle')

    # Save the document
    os.makedirs(os.path.dirname(output_docx_path), exist_ok=True)
    doc.save(output_docx_path)
    print(f"Cover letter Word document saved to {output_docx_path}")


if __name__ == '__main__':
    sample_pdf_path = os.path.join("Data", 'CV_GPT_rev.pdf')
    text = load_and_extract_text(sample_pdf_path)
    output_pdf =  os.path.join("Output", 'test.pdf')
    title = "CV"
    name = extract_applicant_name(load_and_extract_text(sample_pdf_path))
    #create_pdf(output_pdf, title name, text)
    create_pdf(output_pdf,name, text)
    output_docx = os.path.join("Output", 'test_docx.docx')
    """
    (
    output_docx_path,
    sender_name,
    sender_contact,
    recipient_info,
    cover_letter,
    salutation='Dear Hiring Manager,',
    closing='Best regards,
    """
    text = remove_control_characters(text)
    create_docx(output_docx, name,  text)
    text =  "Dear Hiring Manager,\n\nI am writing to express my interest in the AI Developer position. With a Ph.D. in Mathematics and over 3 years of experience in developing AI and machine learning solutions, I have successfully automated invoice reconciliation and standardized financial data formats. My expertise in Python, TensorFlow, PyTorch, scikit-learn, and NLP techniques, combined with my experience in finance, aligns well with your requirements. I am confident in my ability to contribute effectively to your team.\n\nBest regards,\nItay Ben-Dan"
    output_cover_letter = os.path.join("Output", 'test_cover_letter_docx.docx')
    create_cover_letter(text,output_cover_letter)