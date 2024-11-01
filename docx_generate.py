from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re,os

def extract_cv_sections(cv_content):
    """
    Extracts various sections from the CV content by identifying common section headers.

    Args:
        cv_content (str): The full text of the CV.

    Returns:
        dict: A dictionary with section names as keys and section content as values.
    """
    section_patterns = {
        "Name": r"^\s*(.*?)\s*\n",
        "Contact": r"Contact:(.*?)\n\n",
        "Summary": r"Summary:(.*?)\n\n",
        "Experience": r"Experience:(.*?)\n\n",
        "Education": r"Education:(.*?)\n\n",
        "Skills": r"Skills:(.*?)\n\n",
        "Projects": r"Projects:(.*?)\n\n",
        "Publications": r"Publications:(.*?)\n\n",
        "LinkedIn": r"LinkedIn:(.*?)\n",
        "GitHub": r"GitHub:(.*?)\n"
    }

    sections = {}
    for section, pattern in section_patterns.items():
        match = re.search(pattern, cv_content, re.DOTALL | re.MULTILINE)
        if match:
            sections[section] = match.group(1).strip()
        else:
            sections[section] = ""
    return sections
def save_cv_sections_to_file(sections, file_path):
    """
    Saves the extracted CV sections to a file in a structured format.

    Args:
        sections (dict): A dictionary containing CV sections as keys and content as values.
        file_path (str): The path of the file to save the structured output.
    """
    print("In save_cv_sections_to_file")
    with open(file_path, "w", encoding="utf-8") as f:
        for section, content in sections.items():
            f.write(f"{section}:\n")
            f.write(f"{content}\n\n")  # Add spacing for readability
            print(section,":" ,content)
    print(f"CV sections saved to {file_path}")


def load_cv_sections_from_file(file_path):
    """
    Loads CV sections from a structured text file back into a dictionary.

    Args:
        file_path (str): The path of the file containing the CV sections.

    Returns:
        dict: A dictionary with section names as keys and section content as values.
    """
    sections = {}
    current_section = None
    content = []

    with open(file_path, "r", encoding="utf-8") as f:
        for line in f:
            # Check if the line is a section header (e.g., "Name:", "Experience:")
            if line.strip().endswith(":"):
                # Save previous section if there was any
                if current_section:
                    sections[current_section] = "\n".join(content).strip()

                # Start a new section
                current_section = line.strip()[:-1]  # Remove the colon at the end
                content = []
            else:
                # Add lines to the current section content
                content.append(line.strip())

        # Save the last section
        if current_section:
            sections[current_section] = "\n".join(content).strip()

    return sections

def format_header(paragraph, font_size=14, bold=True, color=(31, 56, 100)):
    """
    Formats the paragraph as a header with specified font size, boldness, and color.
    """
    run = paragraph.runs[0]
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(*color)
    run.bold = bold

def add_section(doc, section_title, content, is_bullet=False):
    """
    Adds a section to the document with an optional bullet style.
    """
    # Add section header
    header = doc.add_paragraph(section_title)
    format_header(header)

    # Add section content
    if is_bullet:
        for line in content.split('\n'):
            if line.strip():
                doc.add_paragraph(line, style='List Bullet')
    else:
        for line in content.split('\n'):
            if line.strip():
                paragraph = doc.add_paragraph(line)
                paragraph.paragraph_format.line_spacing = 1.15


def generate_cv_document(file_name, sections):
    """
    Generate and save a CV document based on structured content sections.

    Args:
        file_name (str): Destination path for the generated CV.
        sections (dict): Extracted sections of the CV, e.g., {"Name": ..., "Experience": ...}
    """
    doc = Document()

    # Add Name and Contact Info
    doc.add_paragraph(sections.get("Name", "Name Not Provided"), style="Title").alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_info = sections.get("Contact", "")
    phone = sections.get("Phone", "")
    email = sections.get("Email", "")
    linkedin = sections.get("LinkedIn", "")
    github = sections.get("GitHub", "")

    # Construct contact information string
    contact = f"{contact_info}\n"
    if phone:
        contact += f"Cellular: {phone}\n"
    if email:
        contact += f"Email: {email}\n"
    if linkedin:
        contact += f"LinkedIn: {linkedin}\n"
    if github:
        contact += f"GitHub: {github}\n"

    # Add contact info to the document
    contact_para = doc.add_paragraph(contact)
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add each section with formatting based on content type
    if "Summary" in sections:
        add_section(doc, "Professional Summary", sections["Summary"])

    if "Experience" in sections:
        add_section(doc, "Professional Experience", sections["Experience"], is_bullet=True)

    if "Education" in sections:
        add_section(doc, "Education", sections["Education"])

    if "Skills" in sections:
        add_section(doc, "Key Skills", sections["Skills"], is_bullet=True)

    if "Projects" in sections:
        add_section(doc, "Projects", sections["Projects"], is_bullet=True)

    if "Publications" in sections:
        add_section(doc, "Publications", sections["Publications"])

    # Save the document
    doc.save(file_name)
    print(f"Generated CV document saved as {file_name}")


if __name__ == "__main__":
    # Example of calling the function with sample structured content
    file_name = "Improved_CVNN3.docx"
    finalized_cv_content = {
        "Name": "Itay Ben-Dan",
        "Contact": "Haarava 20, Herzliya, Israel 46100",
        "Phone": "+972544539284",
        "Email": "itaybd@gmail.com",
        "Summary": "Senior Data Scientist and Machine Learning Engineer with extensive experience in data engineering, model development...",
        "Experience": "2017-Present: Machine Learning and Data Science Consultant\n- Developed predictive models from scratch...",
        "Education": "2009: Ph.D. in Mathematics, Technion, Haifa\n2004: M.Sc. in Mathematics, Technion, Haifa...",
        "Skills": "- Programming Languages: Python, C++, R, Java\n- Machine Learning Frameworks: TensorFlow, PyTorch, Scikit-learn...",
        "Projects": "TreeModelVis\n- Developed a visualization tool for decision paths...",
        "Publications": "Published several papers on Computational Geometry and Machine Learning."
    }

    finalized_cv_content = """
    Itay Ben-Dan
    Contact:
    Haarava 20, Herzliya, Israel 46100
    +972544539284
    itaybd@gmail.com
    LinkedIn: Itay Ben-Dan

    Summary:
    Highly skilled AI Developer and Machine Learning Engineer specializing in financial data analytics, invoice reconciliation, and NLP techniques for text analysis...

    Experience:
    2017-Present: Machine Learning and Data Science Consultant
    - Developed and deployed AI models for predictive analysis and automated invoice reconciliation...

    Education:
    2009: Ph.D. in Mathematics, Technion, Haifa
    - Focus: Discrete Geometry

    Skills:
    Python, R, C++, Java
    TensorFlow, PyTorch, Scikit-learn
    Pandas, NumPy, SQL

    Projects:
    TreeModelVis
    - Developed a visualization tool for tree-based models...

    Publications:
    Published numerous papers in Computational Geometry and Combinatorial Theory.
    """

    # Generate the CV

    print(extract_cv_sections(finalized_cv_content))
    generate_cv_document(file_name, extract_cv_sections(finalized_cv_content))
