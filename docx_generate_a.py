import os
from openai import OpenAI
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from data_handling import load_and_extract_text
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from cv_info_extractor import extract_information_from_cv, get_llm_response, parse_extracted_info


# Load OpenAI API key from .env file
load_dotenv('.env', override=True)
openai_api_key = os.getenv('OPENAI_API_KEY')

# Set up OpenAI client
client = OpenAI(api_key=openai_api_key)



def format_header(paragraph, font_size, bold=True, color=(31, 56, 100)):
    """
    Formats the paragraph as a header with specified font size, boldness, and color.
    """
    run = paragraph.runs[0]
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(*color)
    run.bold = bold


def add_plain_text(doc, content):
    """
    Adds plain text content to the document without bullets.
    This will explicitly override any existing bullet styles by using 'Normal' style.
    """
    for line in content.split('\n'):
        line = line.replace('●', '').strip()  # Remove bullet symbols
        if line:
            # Add paragraph with explicit "Normal" style to avoid bullets
            paragraph = doc.add_paragraph(line, style='Normal')
            paragraph.paragraph_format.line_spacing = 1.15  # Line spacing adjustment


def clean_extracted_text(extracted_text):
    """
    Cleans the extracted text from a PDF by removing extra newlines and
    ensuring proper spacing and formatting.
    """
    # Remove excessive newlines and extra spaces
    cleaned_text = "\n".join([line.strip() for line in extracted_text.splitlines() if line.strip()])
    return cleaned_text




def parse_extracted_info(response):
    """
    Parses the structured response from GPT into a dictionary for easy access to the fields.

    Args:
        response (str): The text response from GPT.

    Returns:
        dict: A dictionary containing the extracted information.
    """
    # Initialize an empty dictionary to hold the extracted information
    info_dict = {}

    # Split the response by lines
    lines = response.split("\n")

    # Loop through lines and parse key-value pairs (Name, Email, etc.)
    for line in lines:
        if line.strip():  # Skip empty lines
            key_value = line.split(":")
            if len(key_value) == 2:
                key, value = key_value[0].strip(), key_value[1].strip()
                info_dict[key] = value

    return info_dict


def create_cv_doc(original_cv, job_description, file_name):
    """
    Adapts a CV based on a job description using GPT-4 and saves it as a DOCX file with professional formatting.

    Parameters:
    - original_cv (str): The original CV content.
    - job_description (str): The job description to adapt the CV to.
    - file_name (str): Name of the DOCX file to be saved.
    """
    # Extract personal information using parse_extracted_info
    personal_info = extract_information_from_cv(original_cv)

    name = personal_info.get("Name", "Name Not Found")
    address = personal_info.get("Address", "Address Not Found")
    phone = personal_info.get("Phone", "Phone Not Found")
    email = personal_info.get("Email", "Email Not Found")
    linkedin = personal_info.get("LinkedIn", "")
    github = personal_info.get("GitHub", "")

    prompt = f"""
    Given the following original CV and job description, adapt the CV to align with the job requirements:

    Job Description:
    {job_description}

    Original CV:
    {original_cv}

    Focus on relevant skills, experience, and qualifications for the job.
    """

    # Get the adapted CV from GPT-4
    adapted_cv = get_llm_response(prompt)

    # Create a DOCX file
    doc = Document()

    # Add personal info to the document
    title = doc.add_paragraph(name, style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add address, phone, and email
    contact_paragraph = doc.add_paragraph()
    contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_paragraph.add_run(f"{address}\n")
    contact_paragraph.add_run(f"Cellular: {phone}\n")
    contact_paragraph.add_run(f"Email: {email}\n")

    # Optionally add LinkedIn and GitHub if they are present
    if linkedin:
        contact_paragraph.add_run(f"LinkedIn: {linkedin}\n")
    if github:
        contact_paragraph.add_run(f"GitHub: {github}\n")

    # Ensure contact info is formatted clearly with proper line breaks
    doc.add_paragraph()

    # Split adapted CV into sections and add formatting
    sections = adapted_cv.split('\n\n')

    for section in sections:
        if "Professional Summary" in section:
            parts = section.split('\n', 1)
            content = parts[1] if len(parts) > 1 else parts[0]
            add_formatted_section(doc, 'Professional Summary', content)
        elif "Skills" in section:
            parts = section.split('\n', 1)
            content = parts[1] if len(parts) > 1 else parts[0]
            add_formatted_section(doc, 'Key Skills', content,
                                  bullet_points=True)  # Add bullet points for Skills section
        elif "Experience" in section or "Work Experience" in section:
            parts = section.split('\n', 1)
            content = parts[1] if len(parts) > 1 else parts[0]
            add_formatted_section(doc, 'Professional Experience', content,
                                  bullet_points=True)  # Add bullet points for Experience
        elif "Education" in section:
            parts = section.split('\n', 1)
            content = parts[1] if len(parts) > 1 else parts[0]
            add_formatted_section(doc, 'Education', content, bullet_points=False)  # No bullets for Education

    # Save the document
    doc_path = f'{file_name}.docx'
    doc.save(doc_path)
    return doc_path


def add_horizontal_line(doc):
    """
    Adds a horizontal line to visually separate sections.
    """
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_break()  # Add a break to make space
    p = paragraph._p  # Get the parent XML element of the paragraph
    pPr = p.get_or_add_pPr()  # Get or create the paragraph properties
    pBdr = OxmlElement('w:pBdr')  # Create the border element
    bottom = OxmlElement('w:bottom')  # Create the bottom border element
    bottom.set(qn('w:val'), 'single')  # Single line
    bottom.set(qn('w:sz'), '6')  # Border size
    pBdr.append(bottom)  # Add the bottom border
    pPr.append(pBdr)  # Add the border to the paragraph properties


def add_footer(doc):
    """
    Adds page numbers in the footer.
    """
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.text = "Page "  # Add page number
    run = paragraph.add_run()
    run.font.size = Pt(10)  # Font size for the footer


def add_formatted_section(doc, section_name, content, bullet_points=False):
    """
    Adds a formatted section header and corresponding content to the document without bullets.
    """
    section_paragraph = doc.add_paragraph(section_name)
    format_header(section_paragraph, font_size=14, bold=True)

    # Add the section content without bullets
    add_plain_text(doc, content)


if __name__ == "__main__":
    # Example CV and job description
    sample_pdf_path = os.path.join("Data", 'CV_GPT_rev.pdf')
    original_cv = load_and_extract_text(sample_pdf_path)

    # Clean the extracted text
    original_cv = clean_extracted_text(original_cv)

    job_description = """
    NVIDIA SOC Architecture team is looking for a Senior Data Scientist with SW development skills and HW-System architecture experience. In this position, you will develop datasets, AI models and train AI models for advanced system architecture Power and Performance features, and collaborate with HW & System architects. Strong proficiency in Python, C/C++ required.
    """

    out = os.path.join("Output", "Itay_Ben_Dan_CV_Improved_Visual_2")
    # Create the CV document with professional formatting and visual improvements
    doc_file_path = create_cv_doc(original_cv, job_description, out)

    print(f"The adapted CV has been saved to: {doc_file_path}")
